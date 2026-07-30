"""Turns the AI's generation blocks into real downloadable files."""

import io
import re
import base64
import requests
from urllib.parse import quote

from docx import Document
from fpdf import FPDF
from openpyxl import Workbook
from openpyxl.styles import Font

BLOCK_RE = re.compile(r"```generate:(docx|pdf|excel|image)\s*\n(.*?)```", re.DOTALL)


def _split_title_body(block_text):
    lines = block_text.strip().split("\n", 1)
    title = "Untitled"
    body = block_text.strip()
    if lines and lines[0].lower().startswith("title:"):
        title = lines[0].split(":", 1)[1].strip()
        body = lines[1].strip() if len(lines) > 1 else ""
    return title, body


def make_docx(block_text):
    title, body = _split_title_body(block_text)
    doc = Document()
    doc.add_heading(title, level=1)
    for para in body.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), f"{title[:40] or 'document'}.docx"


def make_pdf(block_text):
    title, body = _split_title_body(block_text)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, title.encode("latin-1", "replace").decode("latin-1"))
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 12)
    for para in body.split("\n\n"):
        if para.strip():
            safe = para.strip().encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(0, 8, safe)
            pdf.ln(2)
    raw = pdf.output()
    return bytes(raw), f"{title[:40] or 'document'}.pdf"


def make_excel(block_text):
    title, body = _split_title_body(block_text)
    wb = Workbook()
    ws = wb.active
    ws.title = title[:31] or "Sheet1"

    rows = [line for line in body.strip().split("\n") if line.strip()]
    for i, line in enumerate(rows):
        cells = [c.strip() for c in line.split(",")]
        ws.append(cells)
        if i == 0:
            for cell in ws[1]:
                cell.font = Font(bold=True)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue(), f"{title[:40] or 'spreadsheet'}.xlsx"


def make_image(block_text):
    prompt = block_text.strip()
    url = f"https://image.pollinations.ai/prompt/{quote(prompt)}"
    resp = requests.get(url, params={"width": 768, "height": 768, "nologo": "true"}, timeout=45)
    resp.raise_for_status()
    safe_name = "".join(c for c in prompt[:30] if c.isalnum() or c in " _-").strip() or "image"
    return resp.content, f"{safe_name}.png"


MAKERS = {
    "docx": (make_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "pdf": (make_pdf, "application/pdf"),
    "excel": (make_excel, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
    "image": (make_image, "image/png"),
}


def extract_generated_files(reply_text):
    """Find every ```generate:type ...``` block in the reply, build the real
    file for each, and return (cleaned_text, [file_dicts]) where cleaned_text
    has the raw blocks replaced with a short human-readable note."""
    files = []

    def _replace(match):
        kind, content = match.group(1), match.group(2)
        maker, mime = MAKERS[kind]
        try:
            data, name = maker(content)
            files.append({
                "name": name,
                "mime": mime,
                "data_b64": base64.b64encode(data).decode("ascii"),
            })
            return f"*(generated {name})*"
        except Exception as e:
            return f"*(couldn't generate {kind}: {e})*"

    cleaned = BLOCK_RE.sub(_replace, reply_text)
    return cleaned, files