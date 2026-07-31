from pypdf import PdfReader
import streamlit as st
import streamlit.components.v1 as components
import os
import sys
import subprocess
import base64
import json
from datetime import datetime
import docx
from dotenv import load_dotenv

from prompts import SYSTEM_PROMPT
from providers import PROVIDERS, has_key, stream_reply
from auth import (
    require_login, current_user, sign_out, display_name, update_display_name
)
from generate import extract_generated_files
import db

load_dotenv()

IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".webp", ".gif")
MAX_DOC_CHARS = 150_000

# Change this if your Institute lives at a different address.
INSTITUTE_URL = "https://alpha-dao-alpha.vercel.app"
INSTITUTE_NAME = "Alpha Institute"

# Alpha Desktop can only be launched when this app is running on the same
# machine as the person using it. Render sets a RENDER env var, so its
# absence means we're running locally.
IS_LOCAL = os.getenv("RENDER") is None
DESKTOP_SCRIPT = "alpha_desktop.py"

DEFAULT_ACCENT = "#7D1B5D"

PALETTES = {
    "Light": {
        "bg": "#FFFFFF",
        "sidebar": "#F0E4EC",
        "sidebar_text": "#3D2A36",
        "sidebar_hover": "#E4D2DE",
        "sidebar_border": "#DCC8D6",
        "text": "#1A1A1A",
        "muted": "#8A7A85",
        "ai_bubble": "#F5EAF1",
        "ai_text": "#3D2A36",
        "border": "#E4D5DE",
        "hover": "#FBF3F8",
    },
    "Dark": {
        "bg": "#16121A",
        "sidebar": "#241C29",
        "sidebar_text": "#EDE2EA",
        "sidebar_hover": "#332738",
        "sidebar_border": "#3A2E42",
        "text": "#ECE7EA",
        "muted": "#9A8F96",
        "ai_bubble": "#2A2230",
        "ai_text": "#ECE7EA",
        "border": "#3A3040",
        "hover": "#2A2230",
    },
}

SUGGESTIONS = [
    ("✍️", "Write", "Help me write something."),
    ("📚", "Learn", "Explain a concept to me in simple terms."),
    ("📄", "Analyze", "Help me analyze a document I'll upload."),
    ("💻", "Code", "Help me write or debug some code."),
]


# ---------- Theme ----------

def active_theme():
    """The palette and accent colour currently in use."""
    mode = st.session_state.get("theme_mode", "Light")
    if mode == "Custom":
        base = st.session_state.get("custom_base", "Light")
        palette = dict(PALETTES[base])
        accent = st.session_state.get("custom_accent", DEFAULT_ACCENT)
    else:
        palette = dict(PALETTES[mode])
        accent = DEFAULT_ACCENT
    palette["accent"] = accent
    return palette


# ---------- Documents ----------

def extract_pdf_text(f):
    reader = PdfReader(f)
    pages = [p.extract_text() for p in reader.pages]
    return "\n\n".join(p for p in pages if p)


def extract_docx_text(f):
    document = docx.Document(f)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def extract_text(f):
    name = f.name.lower()
    if name.endswith(".pdf"):
        return extract_pdf_text(f)
    if name.endswith(".docx"):
        return extract_docx_text(f)
    if name.endswith((".txt", ".md", ".csv")):
        return f.read().decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {f.name}")


def clear_document():
    st.session_state.doc_text = None
    st.session_state.doc_name = None


def clear_image():
    st.session_state.image_bytes = None
    st.session_state.image_name = None


def start_new_chat():
    st.session_state.messages = [{"role": "system", "content": SYSTEM_PROMPT}]
    st.session_state.chat_id = None
    clear_document()
    clear_image()


def derive_title(messages):
    for m in messages:
        if m["role"] == "user":
            return m["content"][:60]
    return "New chat"


def greeting():
    hour = datetime.now().hour
    if hour < 12:
        return "Good morning"
    if hour < 18:
        return "Good afternoon"
    return "Good evening"


@st.cache_data
def logo_base64():
    """Read logo.png once and encode it so it can be embedded directly in HTML."""
    try:
        with open("logo.png", "rb") as f:
            return base64.b64encode(f.read()).decode("utf-8")
    except Exception:
        return None


def render_bubble(role, content):
    """Draw one chat bubble using plain HTML so styling never depends on
    Streamlit's internal (and version-dependent) CSS class names."""
    t = active_theme()
    safe = content.replace("<", "&lt;").replace(">", "&gt;")
    if role == "user":
        st.markdown(
            f"""
            <div style="display:flex; justify-content:flex-end; margin:6px 0;">
                <div style="background:{t['accent']}; color:#fff;
                            border-radius:16px 16px 4px 16px;
                            padding:0.6rem 1rem; max-width:70%;
                            width:fit-content; white-space:pre-wrap;
                            font-size:15px; line-height:1.5;">
                    {safe}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        logo = logo_base64()
        avatar_html = (
            f'<img src="data:image/png;base64,{logo}" style="width:26px;height:26px;border-radius:50%;flex-shrink:0;" />'
            if logo else
            f'<div style="width:26px;height:26px;border-radius:50%;background:{t["accent"]};'
            f'display:flex;align-items:center;justify-content:center;color:#fff;font-size:13px;flex-shrink:0;">a</div>'
        )
        st.markdown(
            f"""
            <div style="display:flex; align-items:flex-end; gap:8px; margin:6px 0;">
                {avatar_html}
                <div style="background:{t['ai_bubble']}; color:{t['ai_text']};
                            border-radius:16px 16px 16px 4px;
                            padding:0.6rem 1rem; max-width:70%;
                            width:fit-content; white-space:pre-wrap;
                            font-size:15px; line-height:1.5;">
                    {safe}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )


def render_files(files, key_prefix):
    """Show a download button (and inline preview for images) for each
    file attached to a message."""
    for i, f in enumerate(files):
        raw = base64.b64decode(f["data_b64"])
        if f["mime"].startswith("image/"):
            st.image(raw, width=300)
        st.download_button(
            f"Download {f['name']}",
            data=raw,
            file_name=f["name"],
            mime=f["mime"],
            key=f"{key_prefix}_{i}",
        )


def render_speak_button(text, key, autoplay=False):
    """A small 'Listen' button using the browser's built-in text-to-speech
    (speechSynthesis) — free, no API, runs entirely client-side.
    When autoplay is on (voice mode), it speaks immediately and then
    restarts listening so the conversation flows hands-free."""
    accent = active_theme()["accent"]
    safe = json.dumps(text)
    auto_js = ""
    if autoplay:
        auto_js = f"""
        window.speechSynthesis.cancel();
        const auto = new SpeechSynthesisUtterance({safe});
        auto.onend = function() {{
            if (!('webkitSpeechRecognition' in window)) return;
            const rec = new webkitSpeechRecognition();
            rec.lang = 'en-US';
            rec.onresult = function(e) {{
                const t = e.results[0][0].transcript;
                const url = new URL(window.parent.location.href);
                url.searchParams.set('voice', t);
                window.parent.location.href = url.toString();
            }};
            rec.start();
        }};
        window.speechSynthesis.speak(auto);
        """
    components.html(
        f"""
        <button id="{key}" style="
            border: 1px solid {accent}44; background: transparent;
            border-radius: 999px; padding: 0.3rem 0.8rem; font-size: 12px;
            cursor: pointer; color: {accent}; font-family: Inter, sans-serif;
            display: inline-flex; align-items: center; gap: 5px;
        ">
        <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="{accent}" stroke-width="2" stroke-linecap="round">
            <path d="M11 5L6 9H2v6h4l5 4V5z"/><path d="M15.5 8.5a5 5 0 010 7"/>
        </svg>
        Listen</button>
        <script>
        document.getElementById("{key}").onclick = function() {{
            window.speechSynthesis.cancel();
            const u = new SpeechSynthesisUtterance({safe});
            window.speechSynthesis.speak(u);
        }};
        {auto_js}
        </script>
        """,
        height=40,
    )


def render_mic_button(key):
    """Mic button using the browser's built-in speech recognition —
    free, no API. Chrome/Edge only. Turns red while listening."""
    accent = active_theme()["accent"]
    components.html(
        f"""
        <button id="{key}" title="Speak your message" style="
            border: none; background: transparent; cursor: pointer;
            padding: 6px; display: flex; align-items: center;
            justify-content: center; border-radius: 50%;
        ">
        <svg id="{key}_icon" width="20" height="20" viewBox="0 0 24 24" fill="none"
             stroke="{accent}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
            <path d="M12 1a3 3 0 00-3 3v8a3 3 0 006 0V4a3 3 0 00-3-3z"/>
            <path d="M19 10v2a7 7 0 01-14 0v-2"/>
            <line x1="12" y1="19" x2="12" y2="23"/>
        </svg>
        </button>
        <script>
        document.getElementById("{key}").onclick = function() {{
            if (!('webkitSpeechRecognition' in window)) {{
                alert('Voice input needs Chrome or Edge.');
                return;
            }}
            const icon = document.getElementById("{key}_icon");
            const recognition = new webkitSpeechRecognition();
            recognition.lang = 'en-US';
            recognition.onstart = function() {{ icon.setAttribute('stroke', '#D9534F'); }};
            recognition.onend = function() {{ icon.setAttribute('stroke', '{accent}'); }};
            recognition.onresult = function(event) {{
                const text = event.results[0][0].transcript;
                const url = new URL(window.parent.location.href);
                url.searchParams.set('voice', text);
                window.parent.location.href = url.toString();
            }};
            recognition.start();
        }};
        </script>
        """,
        height=36,
    )


def launch_desktop_assistant():
    """Open Alpha Desktop in its own console window. Only meaningful when
    this app is running on the same machine as the person using it."""
    if not os.path.exists(DESKTOP_SCRIPT):
        raise FileNotFoundError(
            f"{DESKTOP_SCRIPT} isn't in this folder."
        )
    if os.name == "nt":
        subprocess.Popen(
            [sys.executable, DESKTOP_SCRIPT],
            creationflags=subprocess.CREATE_NEW_CONSOLE,
        )
    else:
        subprocess.Popen([sys.executable, DESKTOP_SCRIPT])


def process_prompt(prompt):
    """Send one user message through the model and save the result.
    Shared by both the real chat input and the suggestion buttons."""
    st.session_state.messages.append({"role": "user", "content": prompt})
    render_bubble("user", prompt)

    payload = list(st.session_state.messages)

    doc = st.session_state.get("doc_text")
    if doc:
        payload.insert(1, {
            "role": "system",
            "content": (
                "The user has attached a document. Answer from it where relevant, "
                "and say clearly when something isn't covered by it.\n\n"
                f"--- DOCUMENT: {st.session_state.doc_name} ---\n{doc}"
            ),
        })

    with st.chat_message("assistant", avatar="logo.png"):
        try:
            reply = st.write_stream(
                stream_reply(
                    st.session_state.provider,
                    payload,
                    st.session_state.get("image_bytes"),
                    st.session_state.get("image_name"),
                )
            )
        except Exception as e:
            reply = f"Error: {e}"
            st.error(reply)

    clean_text, files = extract_generated_files(reply)
    render_speak_button(clean_text, key="speak_live", autoplay=st.session_state.voice_mode)
    if files:
        render_files(files, key_prefix="live_files")

    st.session_state.messages.append({
        "role": "assistant",
        "content": clean_text,
        "files": files,
    })

    st.session_state.chat_id = db.save_chat(
        st.session_state.chat_id,
        user.id,
        derive_title(st.session_state.messages),
        st.session_state.messages,
    )
    db.refresh_caches()


# ---------- Page setup ----------

st.set_page_config(page_title="TheAlpha AI", page_icon="logo.png", layout="centered")

user = require_login()

if "messages" not in st.session_state:
    st.session_state.messages = [{"role": "system", "content": SYSTEM_PROMPT}]
if "chat_id" not in st.session_state:
    st.session_state.chat_id = None
if "provider" not in st.session_state:
    available = [k for k in PROVIDERS if has_key(k)]
    st.session_state.provider = available[0] if available else list(PROVIDERS)[0]
if "pending_prompt" not in st.session_state:
    st.session_state.pending_prompt = None
if "voice_mode" not in st.session_state:
    st.session_state.voice_mode = False
if "theme_mode" not in st.session_state:
    st.session_state.theme_mode = "Light"
if "custom_accent" not in st.session_state:
    st.session_state.custom_accent = DEFAULT_ACCENT
if "custom_base" not in st.session_state:
    st.session_state.custom_base = "Light"

if "voice" in st.query_params:
    st.session_state.pending_prompt = st.query_params.get("voice")
    st.query_params.clear()


# ---------- Make it installable on phones ----------

components.html(
    """
    <script>
    const head = window.parent.document.head;

    function addTag(tag, attrs) {
        for (const [k, v] of Object.entries(attrs)) {
            const existing = head.querySelector(`${tag}[${k}="${v}"]`);
            if (existing) return;
        }
        const el = window.parent.document.createElement(tag);
        for (const [k, v] of Object.entries(attrs)) el.setAttribute(k, v);
        head.appendChild(el);
    }

    addTag('link', {rel: 'manifest', href: '/app/static/manifest.json'});
    addTag('link', {rel: 'apple-touch-icon', href: '/app/static/icon-192.png'});
    addTag('meta', {name: 'apple-mobile-web-app-capable', content: 'yes'});
    addTag('meta', {name: 'apple-mobile-web-app-title', content: 'Alpha'});
    addTag('meta', {name: 'apple-mobile-web-app-status-bar-style', content: 'default'});
    addTag('meta', {name: 'theme-color', content: '#7D1B5D'});
    </script>
    """,
    height=0,
)


# ---------- Fonts + themed styling ----------

T = active_theme()

st.markdown(
    f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Fraunces:ital,wght@0,400;0,500;1,400;1,500&family=Inter:wght@400;500;600&display=swap');

    html, body, .stApp, [data-testid="stAppViewContainer"], [data-testid="stSidebar"] {{
        font-family: 'Inter', sans-serif !important;
    }}

    .stApp, [data-testid="stAppViewContainer"] {{
        background-color: {T['bg']} !important;
        color: {T['text']} !important;
    }}
    [data-testid="stSidebar"] {{
        background-color: {T['sidebar']} !important;
        border-right: 1px solid {T['sidebar_border']} !important;
    }}
    [data-testid="stSidebar"] * {{
        color: {T['sidebar_text']};
    }}
    [data-testid="stSidebar"] hr {{
        border-color: {T['sidebar_border']} !important;
    }}
    [data-testid="stSidebar"] input {{
        background-color: {T['bg']} !important;
        color: {T['sidebar_text']} !important;
        border: 1px solid {T['sidebar_border']} !important;
        border-radius: 8px !important;
    }}
    .stApp p, .stApp h1, .stApp h2, .stApp h3,
    .stApp label, .stApp span {{
        color: {T['text']};
    }}
    [data-testid="stBottom"], [data-testid="stBottomBlockContainer"] {{
        background-color: {T['bg']} !important;
    }}
    div[data-testid="stChatInput"] textarea {{
        background-color: {T['sidebar']} !important;
        color: {T['text']} !important;
        border-radius: 24px !important;
    }}

    .alpha-greeting, .alpha-greeting * {{
        font-family: 'Fraunces', serif !important;
    }}

    div[data-testid="stButton"] button {{
        border-radius: 999px !important;
        border: 1px solid {T['border']} !important;
        background-color: transparent !important;
        color: {T['text']} !important;
        padding: 0.6rem 0.5rem !important;
        font-weight: 500 !important;
        transition: background-color 0.15s ease, border-color 0.15s ease;
    }}
    div[data-testid="stButton"] button:hover {{
        background-color: {T['hover']} !important;
        border-color: {T['accent']} !important;
    }}

    /* Sidebar: flat, quiet list rows instead of bordered pills */
    [data-testid="stSidebar"] div[data-testid="stButton"] button {{
        border: none !important;
        background: transparent !important;
        text-align: left !important;
        font-weight: 400 !important;
        font-size: 14px !important;
        padding: 0.4rem 0.6rem !important;
        border-radius: 8px !important;
        white-space: nowrap !important;
        overflow: hidden !important;
        text-overflow: ellipsis !important;
        display: block !important;
        color: {T['sidebar_text']} !important;
    }}
    [data-testid="stSidebar"] div[data-testid="stButton"] button:hover {{
        background-color: {T['sidebar_hover']} !important;
    }}
    [data-testid="stSidebar"] [data-testid="column"]:last-child div[data-testid="stButton"] button {{
        text-align: center !important;
        padding: 0.4rem !important;
        color: {T['muted']} !important;
    }}
    [data-testid="stSidebar"] [data-testid="column"]:last-child div[data-testid="stButton"] button:hover {{
        color: {T['accent']} !important;
        background-color: {T['sidebar_hover']} !important;
    }}
    /* Institute link button stands out slightly */
    [data-testid="stSidebar"] a[data-testid="stBaseLinkButton-secondary"],
    [data-testid="stSidebar"] div[data-testid="stLinkButton"] a {{
        background-color: {T['accent']} !important;
        border: none !important;
        border-radius: 999px !important;
        color: #fff !important;
        font-size: 14px !important;
        font-weight: 500 !important;
    }}
    [data-testid="stSidebar"] div[data-testid="stLinkButton"] a * {{
        color: #fff !important;
    }}

    /* Pull the mic column so it visually sits beside the send arrow */
    div[data-testid="stBottom"] [data-testid="column"]:nth-of-type(2) {{
        margin-left: -70px !important;
        margin-top: -46px !important;
        z-index: 10;
        position: relative;
        width: 40px !important;
        flex: 0 0 40px !important;
    }}
    div[data-testid="stBottom"] [data-testid="column"]:nth-of-type(2) iframe {{
        width: 40px !important;
    }}

    div[data-testid="stBottomBlockContainer"] {{ padding-bottom: 0.5rem; }}
    div[data-testid="stBottom"] div[data-testid="stVerticalBlock"] {{ gap: 0.25rem; }}
    </style>
    """,
    unsafe_allow_html=True,
)


# ---------- Sidebar ----------

with st.sidebar:
    st.image("logo.png", width=80)
    st.markdown("### TheAlpha AI")
    st.caption(user.email)

    if st.button("New chat", use_container_width=True):
        start_new_chat()
        st.rerun()

    st.divider()

    search = st.text_input(
        "Search chats",
        placeholder="Search your chats...",
        label_visibility="collapsed",
        key="chat_search",
    )

    chats = db.list_chats(user.id)
    if search:
        chats = [c for c in chats if search.lower() in (c.get("title") or "").lower()]
        if not chats:
            st.caption("No chats match that search.")

    for chat in chats:
        row1, row2 = st.columns([5, 1])
        label = chat.get("title") or "Untitled"
        if row1.button(label, key=f"open_{chat['id']}", use_container_width=True):
            messages = db.load_chat(chat["id"])
            if messages:
                st.session_state.messages = messages
                st.session_state.chat_id = chat["id"]
                clear_document()
                clear_image()
                st.rerun()
        if row2.button("🗑", key=f"del_{chat['id']}"):
            db.delete_chat(chat["id"])
            db.refresh_caches()
            if st.session_state.chat_id == chat["id"]:
                start_new_chat()
            st.rerun()

    st.divider()

    st.link_button(
        f"Open {INSTITUTE_NAME}",
        INSTITUTE_URL,
        use_container_width=True,
    )

    with st.expander("Alpha Desktop"):
        st.caption(
            "Controls this computer by voice — opens apps, types, "
            "takes screenshots, sleeps or locks the machine."
        )
        premium = db.is_premium(user.id)
        if premium:
            st.success("Premium — you have access.")
        else:
            st.warning("Premium feature. Contact us to upgrade.")

        if IS_LOCAL and premium:
            if st.button("Launch Alpha Desktop", use_container_width=True):
                try:
                    launch_desktop_assistant()
                    st.success("Opened in a new window.")
                except Exception as e:
                    st.error(f"Couldn't start it: {e}")
        elif not IS_LOCAL and premium:
            st.info(
                "Alpha Desktop runs on your own computer, so it can't be "
                "started from the web. Download the project and run "
                f"`python {DESKTOP_SCRIPT}` on your machine."
            )

    with st.expander("Settings"):
        new_name = st.text_input("Your name", value=display_name(user), key="set_name")
        if st.button("Save name", use_container_width=True):
            try:
                if update_display_name(new_name):
                    st.success("Name updated.")
                    st.rerun()
            except Exception as e:
                st.error(f"Couldn't update name: {e}")

        model_options = [k for k in PROVIDERS if has_key(k)] or list(PROVIDERS.keys())
        st.session_state.provider = st.selectbox(
            "Default model",
            model_options,
            index=model_options.index(st.session_state.provider)
            if st.session_state.provider in model_options else 0,
            key="set_model",
        )

        st.session_state.theme_mode = st.radio(
            "Theme",
            ["Light", "Dark", "Custom"],
            index=["Light", "Dark", "Custom"].index(st.session_state.theme_mode),
            key="set_theme",
        )
        if st.session_state.theme_mode == "Custom":
            st.session_state.custom_base = st.radio(
                "Base", ["Light", "Dark"],
                index=["Light", "Dark"].index(st.session_state.custom_base),
                horizontal=True,
                key="set_custom_base",
            )
            st.session_state.custom_accent = st.color_picker(
                "Accent colour",
                value=st.session_state.custom_accent,
                key="set_accent",
            )

    with st.expander("Usage"):
        stats = db.usage_stats(user.id)
        st.metric("Chats", stats["chats"])
        st.metric("Messages", stats["messages"])

    st.divider()
    if st.button("Sign out", use_container_width=True):
        sign_out()


# ---------- Main area ----------

is_empty_chat = len(st.session_state.messages) <= 1

if is_empty_chat:
    logo = logo_base64()
    logo_html = (
        f'<img src="data:image/png;base64,{logo}" style="width:48px;height:48px;border-radius:50%;margin-bottom:1rem;" />'
        if logo else
        f'<div style="width:48px;height:48px;border-radius:50%;background:{T["accent"]};'
        f'display:flex;align-items:center;justify-content:center;color:#fff;font-size:22px;margin:0 auto 1rem;">a</div>'
    )
    st.markdown(
        f"""
        <div style="text-align:center; padding: 3.5rem 0 2rem;">
            {logo_html}
            <h1 class="alpha-greeting" style="font-size:34px; font-weight:400; font-style:italic; margin:0; color:{T['accent']};">
                {greeting()}, {display_name(user)}
            </h1>
            <p style="color:{T['muted']}; font-size:15px; margin-top:0.5rem;">
                What would you like to do today?
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    cols = st.columns(len(SUGGESTIONS))
    for col, (icon, label, prompt) in zip(cols, SUGGESTIONS):
        with col:
            if st.button(f"{icon}  {label}", use_container_width=True, key=f"suggest_{label}"):
                st.session_state.pending_prompt = prompt
                st.rerun()

else:
    head1, head2 = st.columns([1, 6])
    with head1:
        st.image("logo.png", width=44)
    with head2:
        st.markdown("## TheAlpha AI")
        st.caption("Your personal AI assistant.")

    history = st.session_state.messages[1:]
    last_index = len(history) - 1
    for i, msg in enumerate(history):
        render_bubble(msg["role"], msg["content"])
        # Each Listen button is an embedded frame. Rendering one per reply
        # makes long chats crawl, so only the latest reply gets one.
        if msg["role"] == "assistant" and i == last_index:
            render_speak_button(msg["content"], key=f"speak_hist_{i}")
        if msg["role"] == "assistant" and msg.get("files"):
            render_files(msg["files"], key_prefix=f"hist_{i}")


# ---------- Bottom bar ----------

with st.bottom:
    attached = st.session_state.get("doc_name") or st.session_state.get("image_name")
    if attached:
        acol1, acol2 = st.columns([6, 1])
        icon = "🖼️" if st.session_state.get("image_name") else "📄"
        acol1.caption(f"{icon} {attached}")
        if acol2.button("Remove", use_container_width=True):
            clear_document()
            clear_image()
            st.rerun()

    input_col, mic_col = st.columns([9, 1])
    with input_col:
        user_input = st.chat_input(
            "Message TheAlpha AI...",
            accept_file=True,
            file_type=["pdf", "docx", "txt", "md", "csv", "png", "jpg", "jpeg", "webp", "gif"],
        )
    with mic_col:
        render_mic_button(key="mic_input")

    vcol, scol = st.columns([1, 3])
    with vcol:
        st.session_state.voice_mode = st.toggle(
            "Voice mode",
            value=st.session_state.voice_mode,
            help="Speaks each reply aloud, then listens for your next message automatically.",
        )
    with scol:
        st.caption(PROVIDERS[st.session_state.provider]["model"])


# ---------- Handle input ----------

if st.session_state.pending_prompt:
    prompt_to_send = st.session_state.pending_prompt
    st.session_state.pending_prompt = None
    process_prompt(prompt_to_send)
    st.rerun()

elif user_input:
    if user_input.files:
        f = user_input.files[0]
        if f.name.lower().endswith(IMAGE_EXTS):
            clear_document()
            st.session_state.image_bytes = f.read()
            st.session_state.image_name = f.name
        else:
            clear_image()
            try:
                with st.spinner("Reading document..."):
                    text = extract_text(f)
                if not text.strip():
                    st.warning(
                        f"No readable text in {f.name}. "
                        "If it's a scan, the text isn't extractable."
                    )
                else:
                    st.session_state.doc_text = text[:MAX_DOC_CHARS]
                    st.session_state.doc_name = f.name
            except Exception as e:
                st.error(f"Couldn't read {f.name}: {e}")

    if user_input.text:
        process_prompt(user_input.text)

    st.rerun()